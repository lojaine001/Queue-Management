"""
Convert DEVELOPER_DOC.md to a formatted Word document.
Run: python export_doc.py
Output: DEVELOPER_DOC.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)


def add_inline_code(run):
    """Style a run as inline code."""
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)


def parse_inline(para, text):
    """
    Add text to a paragraph, handling **bold**, `code`, and mixed inline markup.
    """
    # Pattern: **bold**, `code`, or plain text
    token_re = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = token_re.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            add_inline_code(run)
        else:
            if part:
                para.add_run(part)


def style_doc(doc):
    """Apply global document styles."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Heading colours
    heading_colors = {
        'Heading 1': (0x1F, 0x49, 0x7D),   # dark navy
        'Heading 2': (0x2E, 0x74, 0xB5),   # medium blue
        'Heading 3': (0x4A, 0x4A, 0x4A),   # dark grey
    }
    for h, (r, g, b) in heading_colors.items():
        s = doc.styles[h]
        s.font.color.rgb = RGBColor(r, g, b)
        s.font.name = 'Calibri'
        if h == 'Heading 1':
            s.font.size = Pt(18)
        elif h == 'Heading 2':
            s.font.size = Pt(14)
        else:
            s.font.size = Pt(12)

    # Code block style (reuse 'No Spacing' and override)
    code_style = doc.styles['No Spacing']
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(8.5)


# ── Main converter ────────────────────────────────────────────────────────────

def convert(md_path: str, docx_path: str):
    text = Path(md_path).read_text(encoding='utf-8')
    lines = text.splitlines()

    doc = Document()
    style_doc(doc)

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Find header separator row index
        header = table_rows[0]
        data_rows = [r for r in table_rows[1:] if not re.match(r'^\|[-| :]+\|$', r.strip())]

        cols_raw = [c.strip() for c in header.strip('|').split('|')]
        n_cols = len(cols_raw)

        tbl = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set column widths
        available = Inches(6.0)
        col_w = available / n_cols
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = col_w

        # Header row
        hdr_row = tbl.rows[0]
        for j, val in enumerate(cols_raw):
            cell = hdr_row.cells[j]
            set_cell_bg(cell, '2E74B5')
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for r_idx, row_text in enumerate(data_rows):
            cells_raw = [c.strip() for c in row_text.strip('|').split('|')]
            row = tbl.rows[r_idx + 1]
            bg = 'F2F7FC' if r_idx % 2 == 0 else 'FFFFFF'
            for j in range(n_cols):
                val = cells_raw[j] if j < len(cells_raw) else ''
                cell = row.cells[j]
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                parse_inline(p, val)
                for run in p.runs:
                    run.font.size = Pt(9.5)

        doc.add_paragraph()  # spacing after table
        table_rows = []
        in_table = False

    def flush_code(lines_buf):
        if not lines_buf:
            return
        # Add a shaded box for the code block
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, 'F6F8FA')
        # Remove cell border appearance by setting border to light grey
        content = '\n'.join(lines_buf)
        p = cell.paragraphs[0]
        run = p.add_run(content)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]

        # ── Code block ───────────────────────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                flush_code(code_lines)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table row ────────────────────────────────────────────────────────
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Horizontal rule ──────────────────────────────────────────────────
        if line.strip() == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,3})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            title = h_match.group(2)
            # Strip markdown link anchors from TOC entries
            title = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', title)
            p = doc.add_heading(level=level)
            p.clear()
            run = p.add_run(title)
            if level == 1:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            run.font.name = 'Calibri'
            run.bold = True
            i += 1
            continue

        # ── Blockquote (>) ───────────────────────────────────────────────────
        bq_match = re.match(r'^>\s*(.*)', line)
        if bq_match:
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Inches(0.3)
            parse_inline(p, bq_match.group(1))
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.italic = True
            i += 1
            continue

        # ── Bullet list (- ) ─────────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            content = bullet_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.15)
            parse_inline(p, content)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            content = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, content)
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────────
        if line.strip() == '':
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        parse_inline(p, line)
        i += 1

    # Flush any remaining table/code
    if in_table:
        flush_table()
    if in_code_block:
        flush_code(code_lines)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    base = Path(__file__).parent
    convert(
        str(base / 'DEVELOPER_DOC.md'),
        str(base / 'DEVELOPER_DOC.docx'),
    )
